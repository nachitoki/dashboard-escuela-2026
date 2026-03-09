import os
import yaml
import docx
import re
from pathlib import Path

try:
    from google import genai
except ImportError:
    print("Por favor instala: pip install google-genai python-docx")
    exit(1)

BASE_DIR = Path(__file__).parent
OUTPUT_DIR = BASE_DIR / "Outputs" / "Generadas_2026"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Contexto Extraído desde NotebookLM (Caché local para evitar llamadas extra al MCP)
CONTEXTO_REGLAS_NOTEBOOKLM = """
[CONTEXTO DE NOTEBOOKLM - PROGRAMAS DE RELIGIÓN Y PATRIMONIO AONIKENK 2026]
EJES CURRICULARES (1° a 8° BÁSICO):
- Naturaleza y cultura: El ser humano cuida la 'casa común'. (Especialmente en Patrimonio: Huerto local, identidad).
- Persona y sociedad: Promover convivencia, dignidad y bien común.
- Espiritualidad: Búsqueda de sentido, mensaje de Jesucristo.
ENFOQUE DIDÁCTICO:
- Aprendizaje Tridimensional: Integrar conocimientos, habilidades y actitudes.
- Pedagogía de Jesús como modelo: Encuentro personal, diálogo abierto.
- Razón religiosa y cordial: Superar lo meramente cognitivo. Sentido de la realidad.
- En Patrimonio: Fomentar el orgullo por la cultura tehuelche/aonikenk, historia local de la escuela y saberes ancestrales de la región.
VALORES: Dignidad humana, Cuidado de la Casa Común, Inclusión, Fraternidad.
"""

class GeneradorAonikenk:
    def __init__(self):
        self.api_key = os.environ.get("GEMINI_API_KEY")
        if not self.api_key:
            print("⚠️ ADVERTENCIA: No se encontró GEMINI_API_KEY. Operando en modo Mock/Simulación.")
            self.client = None
        else:
            self.client = genai.Client(api_key=self.api_key)

        self.cargar_prompts()

    def cargar_prompts(self):
        workflows_dir = BASE_DIR / ".agents" / "workflows"
        with open(workflows_dir / "prompt_maestro_planificaciones.md", "r", encoding="utf-8") as f:
            self.prompt_planificacion = f.read()
        
        with open(workflows_dir / "prompt_maestro_fichas.md", "r", encoding="utf-8") as f:
            self.prompt_fichas = f.read()

        # with open(BASE_DIR / "ROADMAP.yml", "r", encoding="utf-8") as f:
        #     self.roadmap = list(yaml.safe_load_all(f))
            
        with open(BASE_DIR / "registro_clases.yml", "r", encoding="utf-8") as f:
            self.registro = yaml.safe_load(f)["registro_clases"]

        with open(BASE_DIR / "estudiantes.yml", "r", encoding="utf-8") as f:
            self.contexto_cursos = yaml.safe_load(f)["contexto_estudiantes"]

    def obtener_plan_anual(self, asignatura):
        file_map = {
            "Religión": "plan_anual_religion.yml",
            "Patrimonio": "plan_anual_patrimonio.yml"
        }
        with open(BASE_DIR / file_map[asignatura], "r", encoding="utf-8") as f:
            return yaml.safe_load(f)

    def obtener_siguiente_clase(self, asignatura, curso):
        # Buscar en el registro la primera clase "Pendiente" para ese curso/asignatura
        for reg in self.registro:
            if reg["curso"] == curso and reg["asignatura"] == asignatura and reg["estado"] == "Pendiente":
                return reg
        return None

    def obtener_tema_de_plan(self, asignatura, curso, semana):
        plan = self.obtener_plan_anual(asignatura)
        key = "religion_anual" if asignatura == "Religión" else "patrimonio_anual"
        
        # El plan está indexado por curso (ej: 1_basico)
        curso_key = curso.lower().replace("° ", "_").replace("básico", "basico")
        plan_curso = plan[key].get(curso_key)
        
        if not plan_curso:
            print(f"⚠️ No se encontró plan para el curso {curso_key}")
            return "Tema no encontrado", "Objetivo no encontrado"

        for unidad in plan_curso["unidades"]:
            for clase in unidad["semanas"]:
                if clase["semana"] == semana:
                    return clase["tema"], clase["objetivo"]
        return "Tema no encontrado", "Objetivo no encontrado"

    def generar_siguiente_clase(self, asignatura, curso):
        reg = self.obtener_siguiente_clase(asignatura, curso)
        if not reg:
            print(f"✅ No hay clases pendientes para {asignatura} en {curso}.")
            return None

        semana = reg["semana"]
        tema, objetivo_base = self.obtener_tema_de_plan(asignatura, curso, semana)
        
        # Obtener contexto de estudiantes
        curso_id = curso.lower().replace("° ", "_").replace("básico", "basico")
        contexto = self.contexto_cursos.get(curso_id, {})
        perfil_estudiantes = contexto.get("perfil", "Estudiantes estándar.")
        intereses = contexto.get("intereses", [])
        
        print(f"\n🚀 Generando Clase: {asignatura} | {curso} | Semana {semana}")
        print(f"📝 Tema: {tema}")

        system_instructions_plan = f"{CONTEXTO_REGLAS_NOTEBOOKLM}\n{self.prompt_planificacion}"
        system_instructions_ficha = f"{CONTEXTO_REGLAS_NOTEBOOKLM}\n{self.prompt_fichas}"

        prompt_usuario = f"""
        Debes planificar la Clase de la Semana {semana} de {asignatura} para el curso {curso}.
        
        TEMA PROPUESTO EN EL PLAN ANUAL: {tema}
        OBJETIVO BASE DEL PLAN: {objetivo_base}
        
        CONTEXTO ESPECÍFICO DEL GRUPO DE ESTUDIANTES:
        - Perfil: {perfil_estudiantes}
        - Intereses detectados: {', '.join(intereses)}
        
        INDICACIÓN ESPECIAL: Si hubo un feedback previo en el registro que afecte a esta clase, considéralo.
        RECUERDA: Regla del Objetivo (Taxonomía + Qué + Cómo) y adecuar la didáctica al perfil de los niños.
        """

        if not self.client:
            plan_final = f"# {asignatura} - Clase {semana}\n**Curso:** {curso}\n## 🎯 Objetivo de la Clase\n{objetivo_base}"
            ficha_final = f"# 📚 Ficha de Aprendizaje - {curso}\n**Asignatura:** {asignatura}\n\n## 🚀 ¡Manos a la obra!\n_dibuja el tema: {tema}."
        else:
            try:
                plan_response = self.client.models.generate_content(
                    model="gemini-2.0-flash",
                    contents=prompt_usuario,
                    config={"system_instruction": system_instructions_plan, "temperature": 0.3}
                )
                plan_final = plan_response.text
                
                texto_minusculas = plan_final.lower()
                if any(word in texto_minusculas for word in ["guía", "ficha", "fotocopia", "actividad"]):
                    ficha_prompt = f"Usando la planificación y el contexto de estudiantes ({perfil_estudiantes}), crea la Ficha de Trabajo Estudiantil:\n{plan_final}"
                    ficha_response = self.client.models.generate_content(
                        model="gemini-2.0-flash",
                        contents=ficha_prompt,
                        config={"system_instruction": system_instructions_ficha, "temperature": 0.6}
                    )
                    ficha_final = ficha_response.text
                else:
                    ficha_final = None
            except Exception as e:
                print(f"❌ Error en API: {e}")
                return None

        saved_files = self.guardar_archivos(asignatura, curso, semana, plan_final, ficha_final)
        
        # Marcar como hecho en el registro (en memoria)
        reg["estado"] = "Hecho"
        self.guardar_registro()
        
        return saved_files

    def guardar_registro(self):
        with open(BASE_DIR / "registro_clases.yml", "w", encoding="utf-8") as f:
            yaml.dump({"registro_clases": self.registro}, f, allow_unicode=True, sort_keys=False)

    def guardar_archivos(self, asignatura, curso, semana, plan_txt, ficha_txt):
        nombre_base = f"S{semana}_{asignatura[:3].upper()}_{curso.replace(' ', '')}"
        plan_path = OUTPUT_DIR / f"{nombre_base}_PLAN.md"
        
        with open(plan_path, "w", encoding="utf-8") as f:
            f.write(plan_txt)
        
        # Generar DOCX para el plan
        doc = docx.Document()
        doc.add_heading(f"Planificación: {asignatura} - {curso}", 0)
        doc.add_paragraph(plan_txt)
        docx_path = OUTPUT_DIR / f"{nombre_base}_PLAN.docx"
        doc.save(docx_path)
        
        paths = [plan_path, docx_path]

        if ficha_txt:
            ficha_path = OUTPUT_DIR / f"{nombre_base}_FICHA.md"
            with open(ficha_path, "w", encoding="utf-8") as f:
                f.write(ficha_txt)
            
            doc_f = docx.Document()
            doc_f.add_heading(f"Ficha: {asignatura} - {curso}", 0)
            doc_f.add_paragraph(ficha_txt)
            docx_f_path = OUTPUT_DIR / f"{nombre_base}_FICHA.docx"
            doc_f.save(docx_f_path)
            paths.extend([ficha_path, docx_f_path])
            
        return paths

if __name__ == "__main__":
    generador = GeneradorAonikenk()
    
    clases_lunes = [
        ("Religión", "6° Básico"),
        ("Patrimonio", "1° Básico"),
        ("Religión", "4° Básico")
    ]
    
    for asig, curso in clases_lunes:
        generador.generar_siguiente_clase(asig, curso)
